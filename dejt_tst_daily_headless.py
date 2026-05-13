# -*- coding: utf-8 -*-
r"""
dejt_tst_daily — versão HEADLESS (sem GUI, roda em servidor/Linux/VPS)
=======================================================================
DIFERENÇAS em relação ao original:
  - Sem win32com / Word COM (TOC gerado com python-docx puro)
  - Sem Selenium / WhatsApp Web
  - Envio do DOCX por e-mail via Microsoft Graph API
  - Roda em qualquer OS (Linux, Windows, Mac)

Saídas:
  - PDF : <BASE_DIR>/Diario_J_TST_YYYY-MM-DD.pdf
  - DOCX: <BASE_DIR>/Diario_J_TST_com_variaveis_YYYY-MM-DD.docx
  - Log : <BASE_DIR>/logs/dejt_headless.log

Dependências:
    pip install requests beautifulsoup4 lxml pdf2docx python-docx python-dotenv

Variáveis de ambiente (ou arquivo .env na mesma pasta):
    AZURE_TENANT_ID
    AZURE_CLIENT_ID
    AZURE_CLIENT_SECRET
    GRAPH_SENDER_EMAIL      – remetente
    GRAPH_RECIPIENT_EMAIL   – destinatário (padrão: paollosancheztjmg@gmail.com)
    USE_ADOBE_PDF_SERVICES  – "0" para forçar pdf2docx (padrão: tenta Adobe primeiro)
    PDF_SERVICES_CLIENT_ID / PDF_SERVICES_CLIENT_SECRET  – credenciais Adobe (opcional)
"""
from __future__ import annotations

import os
import sys
import re
import time
import shutil
import logging
from pathlib import Path
from datetime import datetime
from urllib.parse import urljoin

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Carrega .env
# ---------------------------------------------------------------------------
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).resolve().parent / ".env")
except ImportError:
    pass

# ---------------------------------------------------------------------------
# CONFIGURAÇÃO
# ---------------------------------------------------------------------------

BASE_DIR = Path(os.getenv("DEJT_BASE_DIR", r"C:\Users\paollo\Downloads\Codigo"))
LOG_DIR = BASE_DIR / "logs"
LOG_FILE = LOG_DIR / "dejt_headless.log"
DEJT_PAGE = "https://diario.jt.jus.br/cadernos/dejt.html"

REMETENTE = os.getenv("GRAPH_SENDER_EMAIL", "suporte@carvalhofurtadoadv.com.br")
DESTINATARIO = os.getenv("GRAPH_RECIPIENT_EMAIL", "paollosancheztjmg@gmail.com")

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/123.0.0.0 Safari/537.36"
)

# ---------------------------------------------------------------------------
# LOGGING
# ---------------------------------------------------------------------------

def setup_logging() -> None:
    BASE_DIR.mkdir(parents=True, exist_ok=True)
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    fmt = "%(asctime)s | %(levelname)s | %(message)s"
    root = logging.getLogger()
    root.setLevel(logging.INFO)
    for h in list(root.handlers):
        root.removeHandler(h)
    root.addHandler(logging.FileHandler(LOG_FILE, encoding="utf-8"))
    root.addHandler(logging.StreamHandler(sys.stdout))
    for h in root.handlers:
        h.setFormatter(logging.Formatter(fmt=fmt, datefmt="%Y-%m-%d %H:%M:%S"))

# ---------------------------------------------------------------------------
# HTTP SESSION
# ---------------------------------------------------------------------------

def build_session() -> requests.Session:
    from requests.adapters import HTTPAdapter
    try:
        from urllib3.util.retry import Retry
        retry = Retry(total=3, backoff_factor=1.0, status_forcelist=(500, 502, 503, 504))
        adapter = HTTPAdapter(max_retries=retry)
    except Exception:
        adapter = HTTPAdapter(max_retries=3)
    sess = requests.Session()
    sess.headers.update({
        "User-Agent": USER_AGENT,
        "Accept": "text/html,application/xhtml+xml,*/*;q=0.8",
        "Accept-Language": "pt-BR,pt;q=0.9",
    })
    sess.mount("http://", adapter)
    sess.mount("https://", adapter)
    return sess

# ---------------------------------------------------------------------------
# ENCONTRAR PDF DO TST
# ---------------------------------------------------------------------------

def _context_mentions_tst(text: str) -> bool:
    t = text.lower()
    return "tst" in t or "tribunal superior do trabalho" in t or "caderno judiciário do tst" in t


def find_tst_pdf_url(html: str) -> tuple[str, str]:
    try:
        soup = BeautifulSoup(html, "lxml")
    except Exception:
        soup = BeautifulSoup(html, "html.parser")

    anchors = soup.find_all("a", string=lambda s: isinstance(s, str) and "baixar" in s.lower())
    if not anchors:
        anchors = soup.select("a[href$='.pdf']")

    candidates = []
    for a in anchors:
        href = a.get("href", "")
        if not href.lower().endswith(".pdf"):
            continue
        context = " ".join(p.get_text(" ", strip=True) for p in a.parents if hasattr(p, "get_text"))[:300]
        score = 0
        if _context_mentions_tst(context):
            score += 10
        if re.search(r"diario_j_\d+\.pdf$", href.lower()):
            score += 3
        if "j_03" in href.lower():
            score += 4
        candidates.append((score, href, context))

    if not candidates:
        return ("https://diario.jt.jus.br/cadernos/Diario_J_03.pdf", "TST (fallback J_03)")

    candidates.sort(key=lambda x: x[0], reverse=True)
    best = candidates[0]
    return urljoin(DEJT_PAGE, best[1]), best[2] or "Caderno do TST"

# ---------------------------------------------------------------------------
# DOWNLOAD PDF
# ---------------------------------------------------------------------------

def download_pdf(sess: requests.Session, pdf_url: str) -> Path:
    today = datetime.now().strftime("%Y-%m-%d")
    target = BASE_DIR / f"Diario_J_TST_{today}.pdf"
    if target.exists():
        hms = datetime.now().strftime("%H%M%S")
        target = BASE_DIR / f"Diario_J_TST_{today}-{hms}.pdf"

    logging.info(f"Baixando PDF: {pdf_url}")
    with sess.get(pdf_url, stream=True, timeout=(10, 60)) as r:
        r.raise_for_status()
        with open(target, "wb") as f:
            shutil.copyfileobj(r.raw, f)

    with open(target, "rb") as f:
        if f.read(5) != b"%PDF-":
            logging.warning("Arquivo baixado não parece PDF válido.")
    logging.info(f"PDF salvo: {target}")
    return target

# ---------------------------------------------------------------------------
# CONVERSÃO PDF → DOCX
# ---------------------------------------------------------------------------

def convert_pdf_to_docx_via_adobe(pdf_path: Path, docx_output: Path) -> None:
    """Converte usando Adobe PDF Services API (requer credenciais)."""
    client_id = (
        os.getenv("PDF_SERVICES_CLIENT_ID") or
        os.getenv("ADOBE_CLIENT_ID") or
        os.getenv("PDFSERVICES_CLIENT_ID")
    )
    client_secret = (
        os.getenv("PDF_SERVICES_CLIENT_SECRET") or
        os.getenv("ADOBE_CLIENT_SECRET") or
        os.getenv("PDFSERVICES_CLIENT_SECRET")
    )
    if not client_id or not client_secret:
        raise RuntimeError("Credenciais Adobe ausentes.")

    from adobe.pdfservices.operation.auth.service_principal_credentials import ServicePrincipalCredentials
    from adobe.pdfservices.operation.pdf_services import PDFServices
    from adobe.pdfservices.operation.pdf_services_media_type import PDFServicesMediaType
    from adobe.pdfservices.operation.pdfjobs.jobs.export_pdf_job import ExportPDFJob
    from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_params import ExportPDFParams
    from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_target_format import ExportPDFTargetFormat
    from adobe.pdfservices.operation.pdfjobs.result.export_pdf_result import ExportPDFResult

    logging.info("Convertendo PDF→DOCX via Adobe PDF Services...")
    credentials = ServicePrincipalCredentials(client_id=client_id, client_secret=client_secret)
    pdf_services = PDFServices(credentials=credentials)

    with open(pdf_path, "rb") as f:
        input_stream = f.read()

    input_asset = pdf_services.upload(input_stream=input_stream, mime_type=PDFServicesMediaType.PDF)
    params = ExportPDFParams(target_format=ExportPDFTargetFormat.DOCX)
    job = ExportPDFJob(input_asset=input_asset, export_pdf_params=params)
    location = pdf_services.submit(job)
    result = pdf_services.get_job_result(location, ExportPDFResult)
    stream = pdf_services.get_content(result.get_result().get_asset())

    with open(docx_output, "wb") as f:
        f.write(stream.get_input_stream())
    logging.info("Conversão Adobe concluída.")


def convert_pdf_to_docx_via_pdf2docx(pdf_path: Path, docx_output: Path) -> None:
    from pdf2docx import Converter
    logging.info("Convertendo PDF→DOCX com pdf2docx...")
    cv = Converter(str(pdf_path))
    try:
        cv.convert(str(docx_output), start=0, end=None)
    finally:
        cv.close()

# ---------------------------------------------------------------------------
# SUMÁRIO COM PYTHON-DOCX PURO (sem Word COM)
# ---------------------------------------------------------------------------

def inserir_sumario_docx_puro(docx_path: Path) -> bool:
    """
    Gera um sumário simples baseado nos headings do documento,
    usando apenas python-docx (sem Word COM, sem Windows).
    O sumário é inserido logo após o parágrafo 'SUMÁRIO' existente,
    ou criado no início do documento se não encontrado.
    """
    import unicodedata

    logging.info("Gerando sumário com python-docx puro...")
    try:
        doc = Document(str(docx_path))

        # Coletar headings
        headings = []
        for i, para in enumerate(doc.paragraphs):
            style_name = (para.style.name or "").lower() if para.style else ""
            texto = para.text.strip()
            if not texto:
                continue
            nivel = None
            if "heading 1" in style_name:
                nivel = 1
            elif "heading 2" in style_name:
                nivel = 2
            elif "heading 3" in style_name:
                nivel = 3
            if nivel:
                headings.append((nivel, texto, i))

        # Se não achou headings por estilo, tenta por tamanho de fonte
        if not headings:
            for i, para in enumerate(doc.paragraphs):
                texto = para.text.strip()
                if not texto or len(texto) > 120:
                    continue
                try:
                    tamanho = para.runs[0].font.size
                    if tamanho and tamanho.pt >= 12:
                        headings.append((1, texto, i))
                except Exception:
                    pass

        # Localizar parágrafo SUMÁRIO
        sumario_idx = None
        for i, para in enumerate(doc.paragraphs):
            norm = unicodedata.normalize("NFD", para.text.strip().upper())
            norm = "".join(c for c in norm if not unicodedata.combining(c))
            if norm in ("SUMARIO", "INDICE"):
                sumario_idx = i
                break

        # Montar linhas do sumário
        linhas = ["SUMÁRIO\n"]
        for nivel, texto, _ in headings:
            indent = "    " * (nivel - 1)
            linhas.append(f"{indent}{texto}")

        if not headings:
            logging.warning("Nenhum heading encontrado para o sumário.")
            return False

        # Inserir sumário: adicionar ao início do documento se não achou marcador
        if sumario_idx is None:
            # Insere os parágrafos do sumário no início
            primeiro = doc.paragraphs[0]
            for linha in reversed(linhas):
                p = OxmlElement("w:p")
                r = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.text = linha
                r.append(t)
                p.append(r)
                primeiro._p.addprevious(p)
        else:
            # Insere após o parágrafo SUMÁRIO existente
            sumario_para = doc.paragraphs[sumario_idx]
            for linha in linhas[1:]:  # pula o título SUMÁRIO pois já existe
                novo = sumario_para.insert_paragraph_after(linha)
                if novo.runs:
                    novo.runs[0].font.size = Pt(10)
                    novo.runs[0].font.name = "Arial"

        doc.save(str(docx_path))
        logging.info(f"Sumário inserido com {len(headings)} entradas.")
        return True

    except Exception as e:
        logging.error(f"Erro ao gerar sumário: {e}")
        return False

# ---------------------------------------------------------------------------
# ENVIO POR E-MAIL VIA MICROSOFT GRAPH (com anexo DOCX)
# ---------------------------------------------------------------------------

def _obter_token_graph() -> str:
    tenant_id = os.getenv("AZURE_TENANT_ID", "")
    client_id = os.getenv("AZURE_CLIENT_ID", "")
    client_secret = os.getenv("AZURE_CLIENT_SECRET", "")
    if not all([tenant_id, client_id, client_secret]):
        raise EnvironmentError(
            "Defina AZURE_TENANT_ID, AZURE_CLIENT_ID e AZURE_CLIENT_SECRET."
        )
    url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
    resp = requests.post(url, data={
        "grant_type": "client_credentials",
        "client_id": client_id,
        "client_secret": client_secret,
        "scope": "https://graph.microsoft.com/.default",
    }, timeout=15)
    resp.raise_for_status()
    return resp.json()["access_token"]


def enviar_docx_por_email(docx_path: Path) -> bool:
    """Envia o DOCX como anexo via Microsoft Graph."""
    logging.info(f"Enviando {docx_path.name} por e-mail...")
    try:
        token = _obter_token_graph()
    except Exception as e:
        logging.error(f"Falha ao obter token Graph: {e}")
        logging.warning("Configure AZURE_TENANT_ID, AZURE_CLIENT_ID, AZURE_CLIENT_SECRET para envio por e-mail.")
        return False

    import base64
    with open(docx_path, "rb") as f:
        conteudo_b64 = base64.b64encode(f.read()).decode()

    hoje = datetime.now().strftime("%d/%m/%Y")
    payload = {
        "message": {
            "subject": f"Diário de Justiça TST – {hoje}",
            "body": {
                "contentType": "HTML",
                "content": (
                    f"<p>Prezados,</p>"
                    f"<p>Segue em anexo o Diário de Justiça do TST de {hoje}, "
                    f"convertido para DOCX com sumário.</p>"
                    f"<p><em>Gerado automaticamente.</em></p>"
                ),
            },
            "toRecipients": [{"emailAddress": {"address": DESTINATARIO}}],
            "attachments": [
                {
                    "@odata.type": "#microsoft.graph.fileAttachment",
                    "name": docx_path.name,
                    "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "contentBytes": conteudo_b64,
                }
            ],
        },
        "saveToSentItems": True,
    }

    url = f"https://graph.microsoft.com/v1.0/users/{REMETENTE}/sendMail"
    resp = requests.post(url, json=payload, headers={
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }, timeout=60)

    if resp.status_code == 202:
        logging.info(f"E-mail enviado para {DESTINATARIO}.")
        return True
    else:
        logging.error(f"Erro ao enviar e-mail: {resp.status_code} – {resp.text[:300]}")
        return False

# ---------------------------------------------------------------------------
# CHAMAR jt_juris_teste_headless.py (se existir)
# ---------------------------------------------------------------------------

def chamar_jt_juris(docx_path: Path, trt_selection: str = "3") -> int:
    import subprocess
    script_dir = Path(__file__).resolve().parent
    # Tenta a versão headless primeiro, depois a original
    for nome in ("jt_juris_teste_headless.py", "jt_juris_teste 1 (1).py"):
        jt_script = script_dir / nome
        if jt_script.exists():
            break
    else:
        logging.error("Script JT não encontrado.")
        return 1

    env = os.environ.copy()
    env["JT_TRT_SELECTION"] = trt_selection
    env["JT_AUTO_MODE"] = "1"
    env["JT_HEADED"] = "0"          # headless
    env["JT_DOCX_PATH"] = str(docx_path.resolve())

    logging.info(f"Chamando JT Juris [{jt_script.name}] TRT_SELECTION={trt_selection}")
    result = subprocess.run([sys.executable, str(jt_script)], env=env,
                            cwd=str(script_dir))
    logging.info(f"JT Juris finalizado – código {result.returncode}")
    return result.returncode

# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main() -> int:
    setup_logging()
    logging.info("=" * 60)
    logging.info("DEJT TST DAILY — headless")
    logging.info("=" * 60)

    try:
        BASE_DIR.mkdir(parents=True, exist_ok=True)
        sess = build_session()

        # 1) Descobrir URL do PDF
        logging.info(f"Buscando página DEJT: {DEJT_PAGE}")
        resp = sess.get(DEJT_PAGE, timeout=(10, 30))
        resp.raise_for_status()
        pdf_url, label = find_tst_pdf_url(resp.text)
        logging.info(f"URL selecionada: {pdf_url} | {label}")

        # 2) Baixar PDF
        pdf_path = download_pdf(sess, pdf_url)

        # 3) Converter para DOCX
        today = datetime.now().strftime("%Y-%m-%d")
        docx_output = BASE_DIR / f"Diario_J_TST_com_variaveis_{today}.docx"
        use_adobe = str(os.getenv("USE_ADOBE_PDF_SERVICES", "1")).lower() not in ("0", "false")

        if use_adobe:
            try:
                convert_pdf_to_docx_via_adobe(pdf_path, docx_output)
            except Exception as e:
                logging.warning(f"Adobe falhou ({e}). Usando pdf2docx...")
                convert_pdf_to_docx_via_pdf2docx(pdf_path, docx_output)
        else:
            convert_pdf_to_docx_via_pdf2docx(pdf_path, docx_output)

        # 4) Inserir sumário com python-docx puro (sem Word COM)
        inserir_sumario_docx_puro(docx_output)
        logging.info(f"DOCX gerado: {docx_output}")

        # 5) Chamar extração JT
        logging.info("=" * 60)
        logging.info("Iniciando extração JT (TRT3 + TRT24)...")
        try:
            jt_exit = chamar_jt_juris(docx_output, trt_selection="3")
            if jt_exit != 0:
                logging.warning(f"JT terminou com código {jt_exit}.")
        except Exception as e:
            logging.error(f"Erro na extração JT: {e}")

        # 6) Enviar por e-mail (substitui WhatsApp Web)
        logging.info("=" * 60)
        logging.info("Enviando DOCX por e-mail...")
        try:
            enviar_docx_por_email(docx_output)
        except Exception as e:
            logging.error(f"Erro ao enviar e-mail: {e}")

        logging.info("Concluído com sucesso.")
        return 0

    except Exception as e:
        logging.exception(f"Falha no processo: {e}")
        return 1


if __name__ == "__main__":
    sys.exit(main())
