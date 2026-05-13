# -*- coding: utf-8 -*-
"""
jt_juris_teste — versão HEADLESS (sem GUI, roda em servidor/Linux/VPS)
=======================================================================
DIFERENÇAS em relação ao original:
  - Sem SeleniumBase / Chrome / browser (usa requests + BeautifulSoup)
  - Sem win32com / Word COM (sumário gerado com python-docx puro)
  - Sem pyperclip / pywintypes (sem área de transferência do Windows)
  - Roda em qualquer OS via cron, GitHub Actions, VPS Linux

IMPORTANTE: O site do JT (pje.jt.jus.br) é uma SPA Angular/React.
  - Esta versão usa a API REST pública do PJe JT quando disponível
  - Para TRT3: https://pje.trt3.jus.br/pje/api/...
  - Para TRT24: https://pje.trt24.jus.br/pje/api/...
  - Fallback: scraping HTML da página de resultados
  - Se a API mudar, ajuste as URLs em TRT_CONFIG abaixo

Dependências:
    pip install requests beautifulsoup4 lxml python-docx python-dotenv

Variáveis de ambiente (ou .env na mesma pasta):
    JT_TRT_SELECTION  – "1"=TRT3, "2"=TRT24, "3"=Ambos (padrão: "3")
    JT_AUTO_MODE      – "1" para rodar sem interação
    JT_DOCX_PATH      – caminho do DOCX de saída (padrão: ver OUTPUT_DOCX)
    JT_MAX_PAGES      – máximo de páginas por TRT (padrão: 3)
"""
from __future__ import annotations

import os
import re
import sys
import json
import time
import logging
import traceback
import unicodedata
from pathlib import Path
from typing import Optional, List
from datetime import datetime

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ---------------------------------------------------------------------------
# Carrega .env
# ---------------------------------------------------------------------------
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).resolve().parent / ".env")
except ImportError:
    pass

# ---------------------------------------------------------------------------
# LOGGING
# ---------------------------------------------------------------------------

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)],
)
logger = logging.getLogger("JT_HEADLESS")

# ---------------------------------------------------------------------------
# CONFIGURAÇÃO
# ---------------------------------------------------------------------------

BASE_DIR = Path(os.getenv("DEJT_BASE_DIR", r"C:\Users\paollo\Downloads\Codigo"))
DEFAULT_DOCX = str(BASE_DIR / "Diario_J_TST_com_variaveis.docx")
OUTPUT_DOCX = os.getenv("JT_DOCX_PATH", DEFAULT_DOCX)
TRT_SELECTION = os.getenv("JT_TRT_SELECTION", "3")
MAX_PAGES = int(os.getenv("JT_MAX_PAGES", "10"))

# Configuração dos TRTs
# URL base da API REST do PJe — ajuste se necessário
TRT_CONFIG = {
    "TRT3": {
        "nome": "TRT 3ª Região",
        "api_base": "https://pje.trt3.jus.br/pje/api/v1",
        "site_base": "https://pje.trt3.jus.br/pje",
        "consulta_url": "https://pje.trt3.jus.br/consultaprocessual/",
    },
    "TRT24": {
        "nome": "TRT 24ª Região",
        "api_base": "https://pje.trt24.jus.br/pje/api/v1",
        "site_base": "https://pje.trt24.jus.br/pje",
        "consulta_url": "https://pje.trt24.jus.br/consultaprocessual/",
    },
}

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/147.0.0.0 Safari/537.36"
)

# ---------------------------------------------------------------------------
# FALCÃO — repositório nacional unificado de jurisprudência da JT
# Para renovar: abra https://jurisprudencia.jt.jus.br/jurisprudencia-nacional/pesquisa
# em um navegador, faça uma busca, copie sessionId/juristkn/JSESSIONID do DevTools.
# ---------------------------------------------------------------------------
FALCAO_BASE = "https://jurisprudencia.jt.jus.br/jurisprudencia-nacional-backend/api/no-auth/pesquisa"
FALCAO_SESSION_ID = os.getenv("FALCAO_SESSION_ID", "_2mqvyki")
FALCAO_JURISTKN = os.getenv("FALCAO_JURISTKN", "b805efee56c730")
FALCAO_JSESSIONID = os.getenv("FALCAO_JSESSIONID", "v31UiRGyHig8fC_9-3jlH6VWh2fpM0s8ArmEJZ2H")

# Turmas a varrer por TRT (TRT3: 1-11, TRT24: 1-2)
TRT_TURMAS = {
    "TRT3": list(range(1, 12)),
    "TRT24": [1, 2],
}

# ---------------------------------------------------------------------------
# HTTP SESSION
# ---------------------------------------------------------------------------

def build_session() -> requests.Session:
    from requests.adapters import HTTPAdapter
    try:
        from urllib3.util.retry import Retry
        retry = Retry(total=3, backoff_factor=1.5, status_forcelist=(500, 502, 503, 504))
        adapter = HTTPAdapter(max_retries=retry)
    except Exception:
        adapter = HTTPAdapter(max_retries=3)
    sess = requests.Session()
    sess.headers.update({
        "User-Agent": USER_AGENT,
        "Accept": "application/json, text/html, */*",
        "Accept-Language": "pt-BR,pt;q=0.9",
    })
    sess.mount("http://", adapter)
    sess.mount("https://", adapter)
    return sess

SESSION = build_session()

# ---------------------------------------------------------------------------
# UTILITÁRIOS DE TEXTO
# ---------------------------------------------------------------------------

def normalizar(texto: str) -> str:
    """Remove acentos e converte para minúsculas."""
    norm = unicodedata.normalize("NFD", texto or "")
    return "".join(c for c in norm if not unicodedata.combining(c)).lower()


def extrair_numero_processo(texto: str) -> Optional[str]:
    """Extrai número de processo no padrão CNJ (NNNNNNN-DD.AAAA.J.TT.OOOO)."""
    m = re.search(r"\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}", texto)
    return m.group(0) if m else None


def extrair_turma(texto: str) -> Optional[str]:
    """Extrai turma do texto (ex: '1ª Turma', '2ª Turma')."""
    m = re.search(r"(\d{1,2})\s*[ªa°]\s*[Tt]urma", texto)
    return f"{int(m.group(1))}ª" if m else None


def extrair_id_bloco(orgao: str) -> Optional[str]:
    """Extrai identificador TRT+Turma do texto do órgão."""
    if not orgao:
        return None
    txt = normalizar(orgao)
    txt = re.sub(r"[^a-z0-9\s]+", " ", txt)

    # TRT + Turma
    m = re.search(r"\btrt\s*(\d{1,2})\b[^\d]*(\d{1,2})\s*[ªa]?\s*turma", txt)
    if m:
        return f"TRT{int(m.group(1))}_{int(m.group(2))}ª"
    # Apenas TRT
    m = re.search(r"\btrt\s*(\d{1,2})\b", txt)
    if m:
        return f"TRT{int(m.group(1))}"
    # Apenas turma
    m = re.search(r"\b(\d{1,2})\s*[ªa]\s*turma\b", txt)
    if m:
        return f"{int(m.group(1))}ª"
    return None


def descricao_por_identificador(ident: Optional[str]) -> str:
    if not ident:
        return "Processo"
    m = re.match(r"TRT(\d+)_(\d+)ª$", ident)
    if m:
        return f"TRT {m.group(1)} - {m.group(2)}ª Turma"
    if re.match(r"^TRT\d+$", ident):
        num = re.search(r"\d+", ident).group()
        return f"TRT {num} - Acórdãos"
    if re.match(r"^\d+ª$", ident):
        return f"Acórdão {ident} Turma"
    return "Processo"

# ---------------------------------------------------------------------------
# EXTRAÇÃO DE ACÓRDÃOS VIA API REST / SCRAPING
# ---------------------------------------------------------------------------

def _falcao_headers() -> dict:
    return {
        "User-Agent": USER_AGENT,
        "Accept": "application/json, text/plain, */*",
        "Accept-Language": "pt-BR,pt;q=0.9",
        "Referer": "https://jurisprudencia.jt.jus.br/jurisprudencia-nacional/pesquisa",
        "sec-ch-ua-platform": '"Windows"',
        "sec-fetch-dest": "empty",
        "sec-fetch-mode": "cors",
        "sec-fetch-site": "same-origin",
    }


def _falcao_cookies() -> dict:
    return {
        "JSESSIONID": FALCAO_JSESSIONID,
        "SESSION_ID_COOKIE_PUJ": FALCAO_SESSION_ID,
    }


def _limpar_html(texto: str) -> str:
    if not texto:
        return ""
    import html
    txt = re.sub(r"<[^>]+>", " ", texto)
    txt = html.unescape(txt)        # &Ccedil; -> Ç, &nbsp; -> espaço, etc.
    txt = txt.replace("\xa0", " ")  # nbsp residual
    txt = re.sub(r"\s+", " ", txt)
    return txt.strip()


def buscar_acordaos_falcao(trt_key: str, n_turma: Optional[int] = None,
                           max_pages: int = 3) -> List[dict]:
    """
    Busca acórdãos COM EMENTA no Falcão (jurisprudencia.jt.jus.br) para um
    TRT (e opcionalmente Turma específica). Se n_turma for None, não filtra
    por turma. Retorna lista de dicts no formato do gerar_docx.
    """
    # TRT3 usa "01ª Turma" (zero-padded), TRT24 usa "1ª Turma" (sem zero)
    if n_turma is not None:
        orgao_padded = f"{n_turma:02d}ª Turma"
        orgao_simple = f"{n_turma}ª Turma"
    else:
        orgao_padded = ""
        orgao_simple = ""
    headers = _falcao_headers()
    cookies = _falcao_cookies()
    acordaos: List[dict] = []

    # Tenta zero-padded ("01ª Turma") primeiro; se retornar 0, tenta simples ("1ª Turma")
    orgao = orgao_padded
    for page in range(max_pages):
        params = {
            "sessionId": FALCAO_SESSION_ID,
            "latitude": 0,
            "longitude": 0,
            "juristkn": FALCAO_JURISTKN,
            "texto": "",
            "verTodosPrecedentes": "false",
            "tribunais": trt_key,
            "orgaoJulgador": orgao,
            "temEmenta": "S",
            "pesquisaSomenteNasEmentas": "false",
            "colecao": "acordaos",
            "page": page,
            "size": 5,  # site só aceita size=5
        }
        try:
            r = SESSION.get(FALCAO_BASE, params=params, headers=headers,
                            cookies=cookies, timeout=45)
        except Exception as e:
            logger.error(f"  {trt_key} {orgao} pg{page}: erro de rede {e}")
            break

        if r.status_code != 200:
            logger.warning(f"  {trt_key} {orgao} pg{page}: HTTP {r.status_code} "
                           f"({r.text[:140]})")
            if r.status_code in (401, 403) and page == 0:
                logger.warning("  Sessão Falcão expirada — atualize "
                               "FALCAO_SESSION_ID, FALCAO_JURISTKN e "
                               "FALCAO_JSESSIONID no .env.")
            break

        try:
            data = r.json()
        except Exception:
            logger.warning(f"  {trt_key} {orgao} pg{page}: resposta não-JSON")
            break

        docs = data.get("documentos") or []
        if not docs:
            # Primeira página vazia com zero-padded → tenta formato simples
            if page == 0 and orgao == orgao_padded and orgao_simple != orgao_padded:
                logger.info(f"  {trt_key} '{orgao}' sem resultados, tentando '{orgao_simple}'...")
                orgao = orgao_simple
                continue
            break

        for d in docs:
            id_doc = d.get("idDocumentoAcordao", "")
            acordaos.append({
                "numero": d.get("numeroProcesso", ""),
                "turma": d.get("turma", "") or orgao,
                "orgao": f"{trt_key} - {d.get('turma', '') or orgao}",
                "data": d.get("dataJulgamento") or d.get("dataJuntada", ""),
                "data_julgamento": d.get("dataJulgamento", ""),
                "data_juntada": d.get("dataJuntada", ""),
                "ementa": _limpar_html(d.get("ementa", "")),
                "relator": d.get("relator", ""),
                "trt": trt_key,
                "id_acordao": id_doc,
                "link": (f"https://jurisprudencia.jt.jus.br/jurisprudencia-nacional/"
                         f"citacao/acordaos/{trt_key}/{id_doc}") if id_doc else "",
            })

        rotulo = orgao or "(todas as turmas)"
        logger.info(f"  {trt_key} {rotulo} pg{page}: +{len(docs)} acórdãos "
                    f"(total: {data.get('quantidadeTotal', 0)})")

        if len(docs) < 5:
            break
        time.sleep(1.5)  # respeita WAF do Falcão

    return acordaos


def buscar_acordaos_trt(trt_key: str, max_pages: int = 3) -> List[dict]:
    """Varre todas as turmas configuradas do TRT no Falcão."""
    logger.info(f"Buscando acórdãos de {trt_key}...")
    todos: List[dict] = []
    for n in TRT_TURMAS.get(trt_key, []):
        ac = buscar_acordaos_falcao(trt_key, n, max_pages=max_pages)
        todos.extend(ac)
        time.sleep(1.0)  # entre turmas
    if not todos:
        logger.warning(f"  {trt_key}: nenhum acórdão obtido (sessão Falcão "
                       f"pode estar expirada).")
    return todos

# ---------------------------------------------------------------------------
# GERAÇÃO DO DOCX
# ---------------------------------------------------------------------------

def _formatar_data(s: str) -> str:
    """ISO 'YYYY-MM-DD[T...]' -> 'DD/MM/YYYY'. Pass-through se não bater."""
    if not s:
        return ""
    s = str(s)[:10]
    m = re.match(r"(\d{4})-(\d{2})-(\d{2})", s)
    return f"{m.group(3)}/{m.group(2)}/{m.group(1)}" if m else s


def sanitizar_bookmark(nome: str) -> str:
    base = re.sub(r"[^A-Za-z0-9_]+", "_", nome or "")
    if not base or not base[0].isalpha():
        base = f"BM_{base}"
    return base[:40]


def inserir_bookmark(para, bookmark_name: str, bid: int) -> None:
    """Insere bookmark Word no parágrafo via XML puro (sem Word COM)."""
    try:
        p = para._p
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(bid))
        bm_start.set(qn("w:name"), bookmark_name)
        p.append(bm_start)
        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(bid))
        p.append(bm_end)
    except Exception:
        pass


def _nome_tribunal_extenso(trt_key: str) -> str:
    m = re.match(r"TRT(\d+)$", trt_key)
    return f"Tribunal Regional do Trabalho da {int(m.group(1))}ª Região" if m else trt_key


def _formatar_turma_label(turma_str: str) -> str:
    """'1ª TURMA' / '01ª Turma' -> '1ª TURMA'."""
    if not turma_str:
        return ""
    m = re.search(r"(\d{1,2})", turma_str)
    return f"{int(m.group(1))}ª TURMA" if m else turma_str


def _add_ementa_with_inline_links(paragrafo, texto: str):
    """Insere texto de ementa com hyperlinks inline — fonte Arial MT 8pt."""
    pf = paragrafo.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    def add_text_run(t: str):
        if not t:
            return
        r = paragrafo.add_run(t)
        r.font.name = 'Arial MT'
        r.font.size = Pt(8)

    def add_hyperlink_run(display: str, url: str):
        r_id = paragrafo.part.relate_to(url, RT.HYPERLINK, is_external=True)
        hl = OxmlElement('w:hyperlink')
        hl.set(qn('r:id'), r_id)
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        # Font
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Arial MT')
        rFonts.set(qn('w:hAnsi'), 'Arial MT')
        rFonts.set(qn('w:cs'), 'Arial MT')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '16')
        # Blue color + underline (explicit, no style dependency)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(rFonts)
        rPr.append(sz)
        rPr.append(color)
        rPr.append(u)
        run.append(rPr)
        t = OxmlElement('w:t')
        t.text = display
        run.append(t)
        hl.append(run)
        paragrafo._p.append(hl)

    pattern = re.compile(r'<(https?://[^>\s]+)>|(https?://\S+)')
    pos = 0
    texto = (texto or '').rstrip() + ' '
    for m_link in pattern.finditer(texto):
        start, end = m_link.start(), m_link.end()
        # Handle newlines in text before link
        pre_text = texto[pos:start]
        for seg_i, seg in enumerate(pre_text.split('\n')):
            if seg_i > 0:
                paragrafo.add_run().add_break()
            add_text_run(seg)
        url = m_link.group(1) or m_link.group(2)
        bracketed = m_link.group(1) is not None
        if bracketed:
            add_text_run('<')
        add_hyperlink_run(url, url)
        if bracketed:
            add_text_run('>')
        pos = end
    # Handle remaining text with newlines
    remaining = texto[pos:]
    for seg_i, seg in enumerate(remaining.split('\n')):
        if seg_i > 0:
            paragrafo.add_run().add_break()
        add_text_run(seg)


def _buscar_sumario_em_documento(doc: Document):
    """Procura parágrafo 'Sumário' no documento. Retorna o parágrafo ou None."""
    for para in doc.paragraphs:
        norm = unicodedata.normalize("NFD", (para.text or "").strip().upper())
        norm = "".join(c for c in norm if not unicodedata.combining(c))
        if norm in ("SUMARIO", "INDICE"):
            return para
    return None


def _prepare_document_with_sumario(doc: Document) -> None:
    """Garante que o documento tenha parágrafo 'Sumário' no final."""
    if _buscar_sumario_em_documento(doc) is None:
        doc.add_paragraph()
        p_sum = doc.add_paragraph("Sumário")
        if p_sum.runs:
            p_sum.runs[0].bold = True
            p_sum.runs[0].font.size = Pt(12)
            p_sum.runs[0].font.name = 'Arial MT'


def _format_header_line(paragrafo, linha: str):
    """Formata linha de cabeçalho — fonte Arial MT 8pt, justificado, espaçamento 1.0.
    Chave antes de ':' em negrito."""
    parts = linha.split(':', 1)
    if len(parts) == 2 and parts[0].strip() and parts[1] is not None:
        key, value = parts
        run_key = paragrafo.add_run(f"{key.strip()}:")
        run_key.bold = True
        run_key.font.name = 'Arial MT'
        run_key.font.size = Pt(8)
        run_val = paragrafo.add_run(f" {value.strip()}")
        run_val.font.name = 'Arial MT'
        run_val.font.size = Pt(8)
    else:
        run = paragrafo.add_run(linha)
        run.font.name = 'Arial MT'
        run.font.size = Pt(8)
        if linha.strip().lower() == 'ementa:':
            run.bold = True

    pf = paragrafo.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _montar_citacao(ac: dict) -> str:
    """
    Monta linha de citação no formato:
      Tribunal Regional do Trabalho da Xª Região (Yª TURMA). Acórdão: NUM.
      Relator(a): NOME. Data de julgamento: DD/MM/YYYY.
      Juntado aos autos em DD/MM/YYYY. Disponível em: <link>
    """
    trt = ac.get("trt", "")
    turma_label = _formatar_turma_label(ac.get("turma", ""))
    numero = ac.get("numero", "")
    relator = ac.get("relator", "")
    data_jul = _formatar_data(ac.get("data_julgamento", ""))
    data_jnt = _formatar_data(ac.get("data_juntada", ""))
    link = ac.get("link", "")

    partes = [f"{_nome_tribunal_extenso(trt)} ({turma_label})"]
    if numero:
        partes.append(f"Acórdão: {numero}")
    if relator:
        partes.append(f"Relator(a): {relator}")
    if data_jul:
        partes.append(f"Data de julgamento: {data_jul}")
    if data_jnt:
        partes.append(f"Juntado aos autos em {data_jnt}")
    citacao = ". ".join(partes) + "."
    if link:
        citacao += f" Disponível em: <{link}>"
    return citacao


def _inserir_bloco_acordao(ac: dict, anchor, doc, proximo_bid, before_anchor: bool = True):
    """Insere um bloco completo de acórdão no DOCX.

    Formato idêntico ao jt_juris_teste 1 (1).py:
      Processo: NUM
      Órgão Judicante: TRT - Turma
      Relator: NOME
      Juntado aos autos em DD/MM/YYYY
      Ementa:
      [ementa completa]
      Tribunal Regional do Trabalho da Xª Região (Yª TURMA). Acórdão: NUM. ... Disponível em: <link>
    """
    numero = ac.get("numero", "")
    relator = ac.get("relator", "")
    ementa = ac.get("ementa", "")
    orgao_txt = ac.get("orgao", "")
    data_jnt = _formatar_data(ac.get("data_juntada", ""))
    citacao = _montar_citacao(ac)

    linha_juntado = f"Juntado aos autos em {data_jnt}" if data_jnt else "Juntado aos autos"

    # Linhas de cabeçalho (espaçamento 1.0)
    header_lines = [
        f"Processo: {numero}",
        f"Órgão Judicante: {orgao_txt}",
        f"Relator: {relator}",
        linha_juntado,
        "Ementa:",
    ]

    # Texto da ementa + citação (espaçamento duplo)
    texto_ementa = f"{ementa}\n{citacao}"

    if before_anchor and anchor is not None:
        anchor.insert_paragraph_before('')
        for linha in header_lines:
            p = anchor.insert_paragraph_before('')
            _format_header_line(p, linha)
        p_em = anchor.insert_paragraph_before('')
        _add_ementa_with_inline_links(p_em, texto_ementa)
        bm_proc = sanitizar_bookmark(f"BM_PROC_{re.sub(r'[^0-9]', '_', numero)}")
        inserir_bookmark(p_em, bm_proc, proximo_bid())
    else:
        doc.add_paragraph('')
        for linha in header_lines:
            p = doc.add_paragraph()
            _format_header_line(p, linha)
        p_em = doc.add_paragraph()
        _add_ementa_with_inline_links(p_em, texto_ementa)
        bm_proc = sanitizar_bookmark(f"BM_PROC_{re.sub(r'[^0-9]', '_', numero)}")
        inserir_bookmark(p_em, bm_proc, proximo_bid())


def gerar_docx(acordaos_por_trt: dict, docx_path: str) -> bool:
    """
    Gera ou atualiza o DOCX com os acórdãos extraídos.
    Formato idêntico ao jt_juris_teste 1 (1).py:
      [ementa completa]
      Tribunal Regional do Trabalho da Xª Região (Yª TURMA). Acórdão: NUM. ...
    Fonte Arial MT 8pt, justificado, espaçamento duplo, hyperlinks inline.
    Cabeçalhos de turma com espaçamento 1.0.
    """
    logger.info(f"Gerando DOCX: {docx_path}")
    try:
        if os.path.exists(docx_path):
            doc = Document(docx_path)
        else:
            doc = Document()

        # Garantir Sumário no documento
        _prepare_document_with_sumario(doc)

        bookmark_counter = [100]
        turma_bookmarks: dict[str, str] = {}

        def proximo_bid() -> int:
            bookmark_counter[0] += 1
            return bookmark_counter[0]

        # Localizar parágrafo Sumário para inserir ANTES dele
        sumario_para = _buscar_sumario_em_documento(doc)

        for trt_key, acordaos in acordaos_por_trt.items():
            if not acordaos:
                continue

            # Agrupar por turma
            por_turma: dict[str, List[dict]] = {}
            for ac in acordaos:
                ident = extrair_id_bloco(ac.get("orgao") or ac.get("turma") or trt_key) or trt_key
                por_turma.setdefault(ident, []).append(ac)

            for ident, lista in sorted(por_turma.items()):
                # Cabeçalho da turma (Heading 1 com bookmark) — antes do Sumário
                if ident not in turma_bookmarks and sumario_para is not None:
                    descricao = descricao_por_identificador(ident)
                    p_head = sumario_para.insert_paragraph_before(descricao)
                    try:
                        p_head.style = 'Heading 1'
                    except Exception:
                        pass
                    p_head.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    bm_name = sanitizar_bookmark(f"BM_TURMA_{ident}")
                    inserir_bookmark(p_head, bm_name, proximo_bid())
                    turma_bookmarks[ident] = bm_name
                    sumario_para.insert_paragraph_before('')

                for ac in lista:
                    _inserir_bloco_acordao(
                        ac, sumario_para, doc, proximo_bid,
                        before_anchor=(sumario_para is not None)
                    )

        doc.save(docx_path)
        logger.info(f"DOCX salvo: {docx_path}")
        logger.info(
            "Nota: abra o DOCX no Word e pressione Ctrl+A -> F9 para "
            "atualizar os campos PAGEREF do sumário com as páginas reais."
        )
        return True

    except Exception as e:
        logger.error(f"Erro ao gerar DOCX: {e}")
        logger.debug(traceback.format_exc())
        return False

# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main() -> int:
    logger.info("=" * 60)
    logger.info("JT JURIS — headless")
    logger.info(f"TRT_SELECTION={TRT_SELECTION} | DOCX={OUTPUT_DOCX}")
    logger.info("=" * 60)

    # Selecionar TRTs
    trts_para_buscar: List[str] = []
    if TRT_SELECTION == "1":
        trts_para_buscar = ["TRT3"]
    elif TRT_SELECTION == "2":
        trts_para_buscar = ["TRT24"]
    else:
        trts_para_buscar = ["TRT3", "TRT24"]

    # Buscar acórdãos
    acordaos_por_trt: dict[str, List[dict]] = {}
    total = 0
    for trt_key in trts_para_buscar:
        acordaos = buscar_acordaos_trt(trt_key, max_pages=MAX_PAGES)
        acordaos_por_trt[trt_key] = acordaos
        total += len(acordaos)
        logger.info(f"  {trt_key}: {len(acordaos)} acórdãos obtidos")

    logger.info(f"Total de acórdãos: {total}")

    if total == 0:
        logger.warning(
            "Nenhum acórdão obtido.\n"
            "Motivo provável: o site do PJe usa JavaScript/Angular e não serve HTML estático.\n"
            "Soluções:\n"
            "  1. Configure credenciais de API se o TRT oferecer acesso autenticado\n"
            "  2. Use a versão original com Selenium em uma VPS Windows\n"
            "  3. Verifique se o endpoint da API mudou em TRT_CONFIG neste script"
        )
        # Gera DOCX vazio para não quebrar o fluxo do dejt_tst_daily
        gerar_docx({}, OUTPUT_DOCX)
        return 1

    # Gerar DOCX
    Path(OUTPUT_DOCX).parent.mkdir(parents=True, exist_ok=True)
    ok = gerar_docx(acordaos_por_trt, OUTPUT_DOCX)
    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main())
